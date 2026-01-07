"""Smart Markdown Editor with Live Preview and Smart Assistant.

A modern, cross-platform desktop markdown editor built with PySide6 (Qt6 for Python).
"""

import sys
import os
import re
import markdown
from PySide6.QtWidgets import (QApplication, QMainWindow, QWidget, QVBoxLayout,
                               QHBoxLayout, QTextEdit, QSplitter, QMenuBar,
                               QMenu, QFileDialog, QMessageBox, QLabel,
                               QGroupBox, QScrollArea, QPushButton, QDockWidget,
                               QDialog, QLineEdit, QCheckBox)
from PySide6.QtCore import Qt, QTimer, QRegularExpression, QSettings
from PySide6.QtGui import QFont, QColor, QPalette, QSyntaxHighlighter, QTextCharFormat, QTextDocument
from PySide6.QtWebEngineWidgets import QWebEngineView
from collections import Counter
from datetime import datetime


try:
    from docx import Document
    from docx.shared import Inches
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    import weasyprint
    WEASYPRINT_AVAILABLE = True
except ImportError:
    WEASYPRINT_AVAILABLE = False

try:
    import html2text
    HTML2TEXT_AVAILABLE = True
except ImportError:
    HTML2TEXT_AVAILABLE = False

try:
    from pygments.formatters import HtmlFormatter
    PYGMENTS_AVAILABLE = True
except ImportError:
    PYGMENTS_AVAILABLE = False


class MarkdownSyntaxHighlighter(QSyntaxHighlighter):
    """Basic Markdown syntax highlighting for the editor."""

    def __init__(self, document, dark_mode: bool = False):
        super().__init__(document)
        self._dark_mode = bool(dark_mode)
        self._rule_formats = []
        self._fence_re = QRegularExpression(r"^\s{0,3}(```|~~~)")
        self._codeblock_format = QTextCharFormat()
        self._build_formats()

    def set_dark_mode(self, dark_mode: bool) -> None:
        dark_mode = bool(dark_mode)
        if dark_mode == self._dark_mode:
            return
        self._dark_mode = dark_mode
        self._build_formats()
        self.rehighlight()

    def _build_formats(self) -> None:
        self._rule_formats = []

        if self._dark_mode:
            heading_color = QColor("#2f81f7")
            muted_color = QColor("#8b949e")
            rule_color = QColor("#30363d")
            code_fg = QColor("#c9d1d9")
            code_bg = QColor("#161b22")
            link_color = QColor("#2f81f7")
            url_color = QColor("#3fb950")
        else:
            heading_color = QColor("#0b4f9c")
            muted_color = QColor("#6a737d")
            rule_color = QColor("#d0d7de")
            code_fg = QColor("#0969da")
            code_bg = QColor("#f6f8fa")
            link_color = QColor("#0969da")
            url_color = QColor("#1a7f37")

        heading_format = QTextCharFormat()
        heading_format.setForeground(heading_color)
        heading_format.setFontWeight(QFont.Bold)
        self._add_rule(r"^\s{0,3}#{1,6} .*", heading_format)

        blockquote_format = QTextCharFormat()
        blockquote_format.setForeground(muted_color)
        self._add_rule(r"^\s{0,3}>\s.*", blockquote_format)

        list_marker_format = QTextCharFormat()
        list_marker_format.setForeground(muted_color)
        list_marker_format.setFontWeight(QFont.Bold)
        self._add_rule(r"^\s{0,3}([-*+])\s+", list_marker_format)
        self._add_rule(r"^\s{0,3}(\d+)\.\s+", list_marker_format)

        hr_format = QTextCharFormat()
        hr_format.setForeground(rule_color)
        self._add_rule(r"^\s{0,3}(-{3,}|\*{3,}|_{3,})\s*$", hr_format)

        bold_format = QTextCharFormat()
        bold_format.setFontWeight(QFont.Bold)
        self._add_rule(r"\*\*[^\*\n]+\*\*", bold_format)
        self._add_rule(r"__[^_\n]+__", bold_format)

        italic_format = QTextCharFormat()
        italic_format.setFontItalic(True)

        self._add_rule(r"(?<!\*)\*[^\*\n]+\*(?!\*)", italic_format)
        self._add_rule(r"(?<!_)_[^_\n]+_(?!_)", italic_format)

        inline_code_format = QTextCharFormat()
        inline_code_format.setForeground(code_fg)
        inline_code_format.setBackground(code_bg)
        self._add_rule(r"`[^`\n]+`", inline_code_format)

        link_text_format = QTextCharFormat()
        link_text_format.setForeground(link_color)
        self._add_rule(r"\[[^\]]+\](?=\()", link_text_format)

        link_url_format = QTextCharFormat()
        link_url_format.setForeground(url_color)
        self._add_rule(r"\([^\)\s]+\)", link_url_format)

        self._codeblock_format = QTextCharFormat()
        self._codeblock_format.setForeground(code_fg if self._dark_mode else QColor("#24292f"))
        self._codeblock_format.setBackground(code_bg)
        self._codeblock_format.setFontFamily("SF Mono")

    def _add_rule(self, pattern: str, fmt: QTextCharFormat) -> None:
        self._rule_formats.append((QRegularExpression(pattern), fmt))

    def highlightBlock(self, text: str) -> None:
        in_code_block = self.previousBlockState() == 1

        fence_match = self._fence_re.match(text)
        is_fence_line = fence_match.hasMatch()

        if in_code_block:
            self.setFormat(0, len(text), self._codeblock_format)
            if is_fence_line:
                self.setCurrentBlockState(0)
            else:
                self.setCurrentBlockState(1)
            return

        if is_fence_line:
            self.setFormat(0, len(text), self._codeblock_format)
            self.setCurrentBlockState(1)
            return

        self.setCurrentBlockState(0)
        for regex, fmt in self._rule_formats:
            it = regex.globalMatch(text)
            while it.hasNext():
                match = it.next()
                start = match.capturedStart()
                length = match.capturedLength()
                if length > 0:
                    self.setFormat(start, length, fmt)


class MarkdownAnalyzer:
    """Analyze markdown text and produce structure/quality metrics."""

    def __init__(self, text):
        self.text = text
        self.lines = text.split('\n')

    def analyze(self):
        """Perform analysis and return a metrics dictionary."""

        return {
            'word_count': self._count_words(),
            'char_count': len(self.text),
            'line_count': len(self.lines),
            'reading_time': self._estimate_reading_time(),
            'headings': self._analyze_headings(),
            'links': self._analyze_links(),
            'images': self._count_images(),
            'code_blocks': self._count_code_blocks(),
            'lists': self._count_lists(),
            'blockquotes': self._count_blockquotes(),
            'tables': self._count_tables(),
            'readability_score': self._calculate_readability(),
            'structure_quality': self._analyze_structure_quality(),
            'broken_links': self._detect_potential_issues(),
        }

    def _count_words(self):
        """Count words in the document, excluding code blocks."""


        text_without_code = re.sub(r'```.*?```', '', self.text, flags=re.DOTALL)

        text_without_code = re.sub(r'`[^`]+`', '', text_without_code)

        words = re.findall(r'\b\w+\b', text_without_code)
        return len(words)

    def _estimate_reading_time(self):
        """Estimate reading time in minutes (average 200 words/min)."""

        words = self._count_words()
        return max(1, round(words / 200))

    def _analyze_headings(self):
        """Analyze heading structure (H1-H6 counts)."""

        headings = {'h1': 0, 'h2': 0, 'h3': 0, 'h4': 0, 'h5': 0, 'h6': 0}
        for line in self.lines:
            match = re.match(r'^(#{1,6})\s+(.+)$', line.strip())
            if match:
                level = len(match.group(1))
                headings[f'h{level}'] += 1
        return headings

    def _analyze_links(self):
        """Count markdown links."""


        links = re.findall(r'\[([^\]]+)\]\(([^\)]+)\)', self.text)
        return len(links)

    def _count_images(self):
        """Count markdown images."""


        images = re.findall(r'!\[([^\]]*)\]\(([^\)]+)\)', self.text)
        return len(images)

    def _count_code_blocks(self):
        """Count fenced code blocks."""

        code_blocks = re.findall(r'```', self.text)
        return len(code_blocks) // 2

    def _count_lists(self):
        """Count list items (ordered and unordered)."""

        list_items = 0
        for line in self.lines:
            if re.match(r'^\s*[-*+]\s+', line) or re.match(r'^\s*\d+\.\s+', line):
                list_items += 1
        return list_items

    def _count_blockquotes(self):
        """Count blockquote lines."""

        quotes = sum(1 for line in self.lines if line.strip().startswith('>'))
        return quotes

    def _count_tables(self):
        """Count markdown tables (heuristic based)."""


        in_table = False
        table_count = 0
        for line in self.lines:
            if '|' in line and line.strip().startswith('|'):
                if not in_table:
                    table_count += 1
                    in_table = True
            else:
                if in_table and not line.strip().startswith('|'):
                    in_table = False
        return table_count

    def _calculate_readability(self):
        """Calculate a simple readability score (0-100)."""

        score = 100


        paragraphs = self.text.split('\n\n')
        avg_paragraph_length = sum(len(p.split()) for p in paragraphs) / max(len(paragraphs), 1)
        if avg_paragraph_length > 100:
            score -= 10
        elif avg_paragraph_length > 150:
            score -= 20


        headings = self._analyze_headings()
        if headings['h1'] >= 1 and headings['h2'] > 0:
            score += 10


        if sum(headings.values()) == 0 and self._count_words() > 50:
            score -= 15

        return max(0, min(100, score))

    def _analyze_structure_quality(self):
        """Determine a coarse structure quality rating."""

        headings = self._analyze_headings()
        total_headings = sum(headings.values())

        if total_headings == 0:
            return "No structure"
        elif headings['h1'] > 1:
            return "Multiple H1s detected"
        elif headings['h1'] == 1 and headings['h2'] > 0:
            return "Excellent"
        elif total_headings > 0:
            return "Good"
        else:
            return "Needs improvement"

    def _detect_potential_issues(self):
        """Detect potential issues in the document."""

        issues = []


        empty_links = re.findall(r'\[([^\]]+)\]\(\s*\)', self.text)
        if empty_links:
            issues.append(f"{len(empty_links)} empty link(s)")


        headings_text = []
        for line in self.lines:
            match = re.match(r'^#{1,6}\s+(.+)$', line.strip())
            if match:
                headings_text.append(match.group(1))

        duplicates = [h for h, count in Counter(headings_text).items() if count > 1]
        if duplicates:
            issues.append(f"{len(duplicates)} duplicate heading(s)")


        long_lines = sum(1 for line in self.lines if len(line) > 120 and not line.strip().startswith('|'))
        if long_lines > 5:
            issues.append(f"{long_lines} very long lines")

        return issues if issues else ["No issues detected"]


class MarkdownEditor(QMainWindow):
    """Main application window providing editor, preview, and assistant panels."""
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Markdown Editor - Smart Assistant")
        self.setGeometry(100, 100, 1400, 900)

        self._settings = QSettings("smart-markdown-editor", "MarkdownEditor")
        self._dark_mode = bool(self._settings.value("darkMode", False, type=bool))

        self.current_file = None


        self._custom_preview_css_path = self._settings.value("previewCssPath", "", type=str)
        self._custom_preview_css_cache = ""
        self._custom_preview_css_cache_mtime = None


        self._recent_files = self._load_recent_files()

        self._pygments_css_by_theme = {"light": None, "dark": None}


        self.init_ui()


        self.update_timer = QTimer()
        self.update_timer.setSingleShot(True)
        self.update_timer.timeout.connect(self.update_preview)


        self.analysis_timer = QTimer()
        self.analysis_timer.setSingleShot(True)
        self.analysis_timer.timeout.connect(self.update_analysis)


        self._autosave_timer = QTimer()
        self._autosave_timer.setInterval(int(self._settings.value("autoSaveIntervalMs", 30_000, type=int)))
        self._autosave_timer.timeout.connect(self._autosave_tick)
        self._autosave_timer.start()


        self.update_preview()
        self.update_analysis()

    def init_ui(self):

        central_widget = QWidget()
        self.setCentralWidget(central_widget)


        main_splitter = QSplitter(Qt.Horizontal)


        self.editor = QTextEdit()
        self.editor.setPlaceholderText("Type your markdown here...")
        self.editor.textChanged.connect(self.on_text_changed)


        self._syntax_highlighter = MarkdownSyntaxHighlighter(self.editor.document(), dark_mode=self._dark_mode)


        self.preview = QWebEngineView()


        self.assistant_panel = self.create_assistant_panel()


        main_splitter.addWidget(self.editor)
        main_splitter.addWidget(self.preview)
        main_splitter.addWidget(self.assistant_panel)


        main_splitter.setSizes([420, 630, 350])


        layout = QVBoxLayout()
        layout.addWidget(main_splitter)
        central_widget.setLayout(layout)


        self.create_menu_bar()


        self.apply_theme()

    def create_assistant_panel(self):
        """Create the Smart Markdown Assistant panel."""

        panel = QWidget()
        panel_layout = QVBoxLayout()


        title = QLabel("Smart Assistant")
        title_font = QFont()
        title_font.setPointSize(14)
        title_font.setBold(True)
        title.setFont(title_font)
        panel_layout.addWidget(title)


        stats_group = QGroupBox("Document Statistics")
        stats_layout = QVBoxLayout()

        self.stats_labels = {
            'words': QLabel("Words: 0"),
            'chars': QLabel("Characters: 0"),
            'lines': QLabel("Lines: 0"),
            'reading_time': QLabel("Reading time: 1 min"),
        }

        for label in self.stats_labels.values():
            stats_layout.addWidget(label)

        stats_group.setLayout(stats_layout)
        panel_layout.addWidget(stats_group)


        structure_group = QGroupBox("Document Structure")
        structure_layout = QVBoxLayout()

        self.structure_labels = {
            'headings': QLabel("Headings: 0"),
            'links': QLabel("Links: 0"),
            'images': QLabel("Images: 0"),
            'code_blocks': QLabel("Code blocks: 0"),
            'lists': QLabel("List items: 0"),
            'blockquotes': QLabel("Blockquotes: 0"),
            'tables': QLabel("Tables: 0"),
        }

        for label in self.structure_labels.values():
            structure_layout.addWidget(label)

        structure_group.setLayout(structure_layout)
        panel_layout.addWidget(structure_group)


        quality_group = QGroupBox("Quality Analysis")
        quality_layout = QVBoxLayout()

        self.quality_labels = {
            'readability': QLabel("Readability: --"),
            'structure_quality': QLabel("Structure: --"),
        }

        for label in self.quality_labels.values():
            quality_layout.addWidget(label)

        quality_group.setLayout(quality_layout)
        panel_layout.addWidget(quality_group)


        issues_group = QGroupBox("Potential Issues")
        issues_layout = QVBoxLayout()

        self.issues_label = QLabel("No issues detected")
        self.issues_label.setWordWrap(True)
        issues_layout.addWidget(self.issues_label)

        issues_group.setLayout(issues_layout)
        panel_layout.addWidget(issues_group)


        self.format_button = QPushButton("Auto-Format Document")
        self.format_button.clicked.connect(self.auto_format_document)
        panel_layout.addWidget(self.format_button)


        panel_layout.addStretch()

        panel.setLayout(panel_layout)
        return panel

    def create_menu_bar(self):
        menubar = self.menuBar()


        file_menu = menubar.addMenu("File")


        new_action = file_menu.addAction("New")
        new_action.setShortcut("Ctrl+N")
        new_action.triggered.connect(self.new_file)


        open_action = file_menu.addAction("Open")
        open_action.setShortcut("Ctrl+O")
        open_action.triggered.connect(self.open_file)


        self.recent_files_menu = file_menu.addMenu("Recent Files")
        self._rebuild_recent_files_menu()


        save_action = file_menu.addAction("Save")
        save_action.setShortcut("Ctrl+S")
        save_action.triggered.connect(self.save_file)


        save_as_action = file_menu.addAction("Save As...")
        save_as_action.setShortcut("Ctrl+Shift+S")
        save_as_action.triggered.connect(self.save_file_as)

        file_menu.addSeparator()


        export_menu = file_menu.addMenu("Export As")


        export_md_action = export_menu.addAction("Markdown (.md)")
        export_md_action.triggered.connect(lambda: self.export_file('md'))

        export_txt_action = export_menu.addAction("Plain Text (.txt)")
        export_txt_action.triggered.connect(lambda: self.export_file('txt'))

        export_html_action = export_menu.addAction("HTML (.html)")
        export_html_action.triggered.connect(lambda: self.export_file('html'))

        if DOCX_AVAILABLE:
            export_docx_action = export_menu.addAction("Word Document (.docx)")
            export_docx_action.triggered.connect(lambda: self.export_file('docx'))

        if PDF_AVAILABLE:
            export_pdf_action = export_menu.addAction("PDF Document (.pdf)")
            export_pdf_action.triggered.connect(lambda: self.export_file('pdf'))

        export_rtf_action = export_menu.addAction("Rich Text Format (.rtf)")
        export_rtf_action.triggered.connect(lambda: self.export_file('rtf'))

        export_odt_action = export_menu.addAction("OpenDocument Text (.odt)")
        export_odt_action.triggered.connect(lambda: self.export_file('odt'))

        file_menu.addSeparator()


        exit_action = file_menu.addAction("Exit")
        exit_action.setShortcut("Ctrl+Q")
        exit_action.triggered.connect(self.close)


        edit_menu = menubar.addMenu("Edit")


        undo_action = edit_menu.addAction("Undo")
        undo_action.setShortcut("Ctrl+Z")
        undo_action.triggered.connect(self.editor.undo)


        redo_action = edit_menu.addAction("Redo")
        redo_action.setShortcut("Ctrl+Y")
        redo_action.triggered.connect(self.editor.redo)

        edit_menu.addSeparator()

        find_action = edit_menu.addAction("Find...")
        find_action.setShortcut("Ctrl+F")
        find_action.triggered.connect(self.open_find_dialog)

        replace_action = edit_menu.addAction("Replace...")
        replace_action.setShortcut("Ctrl+H")
        replace_action.triggered.connect(self.open_replace_dialog)


        view_menu = menubar.addMenu("View")

        self.dark_mode_action = view_menu.addAction("Dark Mode")
        self.dark_mode_action.setCheckable(True)
        self.dark_mode_action.setChecked(self._dark_mode)
        self.dark_mode_action.triggered.connect(self.toggle_dark_mode)

        view_menu.addSeparator()

        preview_css_action = view_menu.addAction("Preview CSS...")
        preview_css_action.triggered.connect(self.choose_preview_css)

        clear_preview_css_action = view_menu.addAction("Clear Preview CSS")
        clear_preview_css_action.triggered.connect(self.clear_preview_css)

    def _ensure_find_replace_dialog(self) -> None:
        if hasattr(self, "_find_replace_dialog") and self._find_replace_dialog is not None:
            return

        dlg = QDialog(self)
        dlg.setWindowTitle("Find / Replace")
        dlg.setModal(False)

        layout = QVBoxLayout()

        find_row = QHBoxLayout()
        find_row.addWidget(QLabel("Find:"))
        self._find_input = QLineEdit()
        self._find_input.setPlaceholderText("Text to find")
        find_row.addWidget(self._find_input)
        layout.addLayout(find_row)

        replace_row = QHBoxLayout()
        replace_row.addWidget(QLabel("Replace:"))
        self._replace_input = QLineEdit()
        self._replace_input.setPlaceholderText("Replacement text")
        replace_row.addWidget(self._replace_input)
        layout.addLayout(replace_row)

        options_row = QHBoxLayout()
        self._match_case_cb = QCheckBox("Match case")
        options_row.addWidget(self._match_case_cb)
        options_row.addStretch()
        layout.addLayout(options_row)

        buttons_row = QHBoxLayout()
        self._find_prev_btn = QPushButton("Find Previous")
        self._find_next_btn = QPushButton("Find Next")
        self._replace_btn = QPushButton("Replace")
        self._replace_all_btn = QPushButton("Replace All")
        close_btn = QPushButton("Close")

        self._find_prev_btn.clicked.connect(lambda: self.find_text(backward=True))
        self._find_next_btn.clicked.connect(lambda: self.find_text(backward=False))
        self._replace_btn.clicked.connect(self.replace_one)
        self._replace_all_btn.clicked.connect(self.replace_all)
        close_btn.clicked.connect(dlg.close)

        buttons_row.addWidget(self._find_prev_btn)
        buttons_row.addWidget(self._find_next_btn)
        buttons_row.addWidget(self._replace_btn)
        buttons_row.addWidget(self._replace_all_btn)
        buttons_row.addWidget(close_btn)
        layout.addLayout(buttons_row)

        dlg.setLayout(layout)

        self._find_input.returnPressed.connect(lambda: self.find_text(backward=False))
        self._replace_input.returnPressed.connect(self.replace_one)

        self._find_replace_dialog = dlg
        self._apply_dialog_theme()

    def open_find_dialog(self) -> None:
        self._ensure_find_replace_dialog()
        self._replace_input.setEnabled(False)
        self._replace_btn.setEnabled(False)
        self._replace_all_btn.setEnabled(False)
        self._find_replace_dialog.show()
        self._find_replace_dialog.raise_()
        self._find_replace_dialog.activateWindow()
        self._find_input.setFocus()
        self._find_input.selectAll()

    def open_replace_dialog(self) -> None:
        self._ensure_find_replace_dialog()
        self._replace_input.setEnabled(True)
        self._replace_btn.setEnabled(True)
        self._replace_all_btn.setEnabled(True)
        self._find_replace_dialog.show()
        self._find_replace_dialog.raise_()
        self._find_replace_dialog.activateWindow()
        self._find_input.setFocus()
        self._find_input.selectAll()

    def _text_find_flags(self, backward: bool) -> QTextDocument.FindFlags:
        flags = QTextDocument.FindFlags()
        if backward:
            flags |= QTextDocument.FindBackward
        if getattr(self, "_match_case_cb", None) is not None and self._match_case_cb.isChecked():
            flags |= QTextDocument.FindCaseSensitively
        return flags

    def find_text(self, backward: bool = False) -> bool:
        self._ensure_find_replace_dialog()
        needle = self._find_input.text()
        if not needle:
            return False

        doc = self.editor.document()
        cursor = self.editor.textCursor()
        flags = self._text_find_flags(backward)

        found = doc.find(needle, cursor, flags)
        if found.isNull():

            wrap_cursor = self.editor.textCursor()
            if backward:
                wrap_cursor.movePosition(wrap_cursor.End)
            else:
                wrap_cursor.movePosition(wrap_cursor.Start)

            found = doc.find(needle, wrap_cursor, flags)

        if found.isNull():
            QMessageBox.information(self, "Find", f"'{needle}' not found")
            return False

        self.editor.setTextCursor(found)
        self.editor.ensureCursorVisible()
        return True

    def replace_one(self) -> None:
        self._ensure_find_replace_dialog()
        needle = self._find_input.text()
        if not needle:
            return

        replacement = self._replace_input.text()

        cursor = self.editor.textCursor()
        selected = cursor.selectedText()

        match_case = self._match_case_cb.isChecked()
        matches = False
        if cursor.hasSelection():
            if match_case:
                matches = selected == needle
            else:
                matches = selected.lower() == needle.lower()

        if not matches:
            if not self.find_text(backward=False):
                return
            cursor = self.editor.textCursor()

        cursor.beginEditBlock()
        cursor.insertText(replacement)
        cursor.endEditBlock()
        self.find_text(backward=False)

    def replace_all(self) -> None:
        self._ensure_find_replace_dialog()
        needle = self._find_input.text()
        if not needle:
            return

        replacement = self._replace_input.text()
        doc = self.editor.document()

        flags = QTextDocument.FindFlags()
        if self._match_case_cb.isChecked():
            flags |= QTextDocument.FindCaseSensitively

        cursor = self.editor.textCursor()
        cursor.beginEditBlock()


        scan_cursor = self.editor.textCursor()
        scan_cursor.movePosition(scan_cursor.Start)

        count = 0
        while True:
            found = doc.find(needle, scan_cursor, flags)
            if found.isNull():
                break
            found.insertText(replacement)
            count += 1
            scan_cursor = found

        cursor.endEditBlock()
        QMessageBox.information(self, "Replace All", f"Replaced {count} occurrence(s).")

    def _apply_dialog_theme(self) -> None:
        if not hasattr(self, "_find_replace_dialog") or self._find_replace_dialog is None:
            return
        if self._dark_mode:
            self._find_replace_dialog.setStyleSheet("""
                QDialog { background-color: #0d1117; color: #c9d1d9; }
                QLabel { color: #c9d1d9; }
                QLineEdit { background-color: #161b22; color: #c9d1d9; border: 1px solid #30363d; padding: 4px; }
                QCheckBox { color: #c9d1d9; }
                QPushButton { border: 1px solid #30363d; padding: 6px 10px; }
            """)
        else:
            self._find_replace_dialog.setStyleSheet("")

    def toggle_dark_mode(self, checked: bool):
        self._dark_mode = bool(checked)
        self._settings.setValue("darkMode", self._dark_mode)
        self.apply_theme()
        self.update_preview()

    def apply_theme(self):
        if self._dark_mode:
            editor_bg = "#0d1117"
            editor_fg = "#c9d1d9"
            border = "#30363d"
            selection_bg = "#2f81f7"
            selection_fg = "#ffffff"

            self.editor.setStyleSheet(f"""
                QTextEdit {{
                    background-color: {editor_bg};
                    color: {editor_fg};
                    border: 1px solid {border};
                    font-family: 'SF Mono', 'Monaco', 'Menlo', 'Consolas', monospace;
                    font-size: 14px;
                    padding: 10px;
                    selection-background-color: {selection_bg};
                    selection-color: {selection_fg};
                }}
            """)

            self.assistant_panel.setStyleSheet(f"""
                QWidget {{
                    background-color: {editor_bg};
                    color: {editor_fg};
                }}
                QGroupBox {{
                    border: 1px solid {border};
                    margin-top: 8px;
                    padding: 8px;
                }}
                QGroupBox::title {{
                    subcontrol-origin: margin;
                    left: 10px;
                    padding: 0 4px 0 4px;
                }}
                QPushButton {{
                    border: 1px solid {border};
                    padding: 6px 10px;
                }}
            """)
        else:
            self.editor.setStyleSheet("""
                QTextEdit {
                    background-color: #ffffff;
                    color: #333333;
                    border: 1px solid #ddd;
                    font-family: 'SF Mono', 'Monaco', 'Menlo', 'Consolas', monospace;
                    font-size: 14px;
                    padding: 10px;
                    selection-background-color: #0078d4;
                    selection-color: #ffffff;
                }
            """)

            self.assistant_panel.setStyleSheet("")

        if hasattr(self, "_syntax_highlighter") and self._syntax_highlighter is not None:
            self._syntax_highlighter.set_dark_mode(self._dark_mode)

        self._apply_dialog_theme()

    def on_text_changed(self):


        self.update_timer.start(300)

        self.analysis_timer.start(800)

    def update_preview(self):

        markdown_text = self.editor.toPlainText()


        html = markdown.markdown(markdown_text, extensions=['codehilite', 'tables', 'toc'])

        theme_key = "dark" if self._dark_mode else "light"
        pygments_css = ""
        if PYGMENTS_AVAILABLE:
            if self._pygments_css_by_theme[theme_key] is None:
                style_name = "monokai" if self._dark_mode else "default"
                self._pygments_css_by_theme[theme_key] = HtmlFormatter(style=style_name).get_style_defs('.codehilite')
            pygments_css = self._pygments_css_by_theme[theme_key]

        if self._dark_mode:
            body_bg = "#0d1117"
            body_fg = "#c9d1d9"
            border = "#30363d"
            muted = "#8b949e"
            link = "#2f81f7"
            code_bg = "#161b22"
        else:
            body_bg = "#fff"
            body_fg = "#333"
            border = "#eaecef"
            muted = "#6a737d"
            link = "#0366d6"
            code_bg = "#f6f8fa"


        custom_css = self._get_custom_preview_css()
        styled_html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <style>
                {pygments_css}
                body {{
                    font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
                    line-height: 1.6;
                    color: {body_fg};
                    max-width: 800px;
                    margin: 0 auto;
                    padding: 20px;
                    background-color: {body_bg};
                }}
                h1, h2, h3, h4, h5, h6 {{
                    margin-top: 24px;
                    margin-bottom: 16px;
                    font-weight: 600;
                    line-height: 1.25;
                }}
                h1 {{ font-size: 2em; border-bottom: 1px solid {border}; padding-bottom: 0.3em; }}
                h2 {{ font-size: 1.5em; border-bottom: 1px solid {border}; padding-bottom: 0.3em; }}
                h3 {{ font-size: 1.25em; }}
                h4 {{ font-size: 1em; }}
                h5 {{ font-size: 0.875em; }}
                h6 {{ font-size: 0.85em; color: {muted}; }}
                p {{ margin-bottom: 16px; }}
                code {{
                    background-color: {code_bg};
                    border-radius: 3px;
                    font-size: 85%;
                    margin: 0;
                    padding: 0.2em 0.4em;
                }}
                pre {{
                    background-color: {code_bg};
                    border-radius: 6px;
                    padding: 16px;
                    overflow: auto;
                    font-size: 85%;
                    line-height: 1.45;
                }}
                .codehilite {{
                    background-color: {code_bg};
                    border-radius: 6px;
                    padding: 16px;
                    overflow: auto;
                    margin-bottom: 16px;
                }}
                .codehilite pre {{
                    margin: 0;
                    padding: 0;
                    background: transparent;
                }}
                pre code {{
                    background-color: transparent;
                    border: 0;
                    display: inline;
                    line-height: inherit;
                    margin: 0;
                    max-width: auto;
                    overflow: visible;
                    padding: 0;
                    word-wrap: normal;
                }}
                blockquote {{
                    border-left: 0.25em solid {border};
                    color: {muted};
                    padding: 0 1em;
                    margin: 0 0 16px 0;
                }}
                table {{
                    border-spacing: 0;
                    border-collapse: collapse;
                    margin-bottom: 16px;
                }}
                table th, table td {{
                    border: 1px solid {border};
                    padding: 6px 13px;
                }}
                table th {{
                    background-color: {code_bg};
                    font-weight: 600;
                }}
                table tr:nth-child(2n) {{
                    background-color: {code_bg};
                }}
                ul, ol {{
                    padding-left: 2em;
                    margin-bottom: 16px;
                }}
                li {{
                    margin-bottom: 0.25em;
                }}
                a {{
                    color: {link};
                    text-decoration: none;
                }}
                a:hover {{
                    text-decoration: underline;
                }}
                img {{
                    max-width: 100%;
                    height: auto;
                }}
                hr {{
                    border: none;
                    border-top: 1px solid {border};
                    height: 1px;
                    margin: 24px 0;
                }}
                {custom_css}
            </style>
        </head>
        <body>
            {html}
        </body>
        </html>
        """


        self.preview.setHtml(styled_html)

    def update_analysis(self):
        """Update the Smart Assistant panel with document analysis."""

        markdown_text = self.editor.toPlainText()

        if not markdown_text.strip():

            self.stats_labels['words'].setText("Words: 0")
            self.stats_labels['chars'].setText("Characters: 0")
            self.stats_labels['lines'].setText("Lines: 0")
            self.stats_labels['reading_time'].setText("Reading time: 1 min")

            self.structure_labels['headings'].setText("Headings: 0")
            self.structure_labels['links'].setText("Links: 0")
            self.structure_labels['images'].setText("Images: 0")
            self.structure_labels['code_blocks'].setText("Code blocks: 0")
            self.structure_labels['lists'].setText("List items: 0")
            self.structure_labels['blockquotes'].setText("Blockquotes: 0")
            self.structure_labels['tables'].setText("Tables: 0")

            self.quality_labels['readability'].setText("Readability: --")
            self.quality_labels['structure_quality'].setText("Structure: --")
            self.issues_label.setText("No issues detected")
            return


        analyzer = MarkdownAnalyzer(markdown_text)
        metrics = analyzer.analyze()


        self.stats_labels['words'].setText(f"Words: {metrics['word_count']}")
        self.stats_labels['chars'].setText(f"Characters: {metrics['char_count']}")
        self.stats_labels['lines'].setText(f"Lines: {metrics['line_count']}")
        self.stats_labels['reading_time'].setText(f"Reading time: {metrics['reading_time']} min")


        total_headings = sum(metrics['headings'].values())
        heading_breakdown = ', '.join([f"H{i}: {metrics['headings'][f'h{i}']}"
                                      for i in range(1, 7) if metrics['headings'][f'h{i}'] > 0])
        self.structure_labels['headings'].setText(f"Headings: {total_headings} ({heading_breakdown})" if heading_breakdown else f"Headings: {total_headings}")
        self.structure_labels['links'].setText(f"Links: {metrics['links']}")
        self.structure_labels['images'].setText(f"Images: {metrics['images']}")
        self.structure_labels['code_blocks'].setText(f"Code blocks: {metrics['code_blocks']}")
        self.structure_labels['lists'].setText(f"List items: {metrics['lists']}")
        self.structure_labels['blockquotes'].setText(f"Blockquotes: {metrics['blockquotes']}")
        self.structure_labels['tables'].setText(f"Tables: {metrics['tables']}")


        readability = metrics['readability_score']
        readability_color = "green" if readability >= 80 else "orange" if readability >= 60 else "red"
        self.quality_labels['readability'].setText(f"Readability: {readability}/100")
        self.quality_labels['readability'].setStyleSheet(f"color: {readability_color}; font-weight: bold;")

        structure_quality = metrics['structure_quality']
        structure_color = "green" if structure_quality == "Excellent" else "orange" if structure_quality == "Good" else "red"
        self.quality_labels['structure_quality'].setText(f"Structure: {structure_quality}")
        self.quality_labels['structure_quality'].setStyleSheet(f"color: {structure_color}; font-weight: bold;")


        issues_text = "\n".join(f"• {issue}" for issue in metrics['broken_links'])
        self.issues_label.setText(issues_text)

    def auto_format_document(self):
        """Auto-format the markdown document with best practices."""

        markdown_text = self.editor.toPlainText()

        if not markdown_text.strip():
            QMessageBox.information(self, "Auto-Format", "Document is empty. Nothing to format.")
            return


        formatted_text = self.format_markdown(markdown_text)


        self.editor.setPlainText(formatted_text)

        QMessageBox.information(self, "Auto-Format", "Document formatted successfully!")

    def format_markdown(self, text):
        """Apply auto-formatting rules to markdown text."""

        lines = text.split('\n')
        formatted_lines = []
        prev_was_heading = False
        prev_was_empty = False

        for i, line in enumerate(lines):
            stripped = line.strip()


            if stripped.startswith('#') and i > 0 and not prev_was_empty:
                formatted_lines.append('')


            if stripped.startswith('#'):
                match = re.match(r'^(#{1,6})(\S)', stripped)
                if match:
                    level = match.group(1)
                    rest = stripped[len(level):]
                    formatted_lines.append(f"{level} {rest}")
                    prev_was_heading = True
                    prev_was_empty = False
                    continue


            if re.match(r'^(\s*)([-*+])(\S)', line):
                match = re.match(r'^(\s*)([-*+])(.*)$', line)
                indent = match.group(1)
                marker = match.group(2)
                content = match.group(3).strip()
                formatted_lines.append(f"{indent}{marker} {content}")
                prev_was_heading = False
                prev_was_empty = False
                continue


            if re.match(r'^(\s*)(\d+\.)(\S)', line):
                match = re.match(r'^(\s*)(\d+\.)(.*)$', line)
                indent = match.group(1)
                marker = match.group(2)
                content = match.group(3).strip()
                formatted_lines.append(f"{indent}{marker} {content}")
                prev_was_heading = False
                prev_was_empty = False
                continue


            if not stripped:

                if not prev_was_empty:
                    formatted_lines.append('')
                    prev_was_empty = True
            else:
                formatted_lines.append(line)
                prev_was_empty = False
                prev_was_heading = False


        while formatted_lines and not formatted_lines[-1]:
            formatted_lines.pop()


        return '\n'.join(formatted_lines)

    def new_file(self):
        self.editor.clear()
        self.current_file = None
        self.editor.document().setModified(False)

    def open_file(self):
        file_path, _ = QFileDialog.getOpenFileName(
            self, "Open Markdown File", "", "Markdown Files (*.md);;All Files (*)"
        )

        if file_path:
            self.open_file_path(file_path)

    def open_file_path(self, file_path: str) -> None:
        if not file_path:
            return
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            self.editor.setPlainText(content)
            self.editor.document().setModified(False)
            self.current_file = file_path
            self._add_recent_file(file_path)
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Could not open file: {str(e)}")

    def save_file(self):
        if hasattr(self, 'current_file') and self.current_file:
            self.save_to_file(self.current_file)
        else:
            self.save_file_as()

    def save_file_as(self):
        file_path, _ = QFileDialog.getSaveFileName(
            self, "Save Markdown File", "", "Markdown Files (*.md);;All Files (*)"
        )

        if file_path:
            self.save_to_file(file_path)

    def save_to_file(self, file_path):
        self._save_to_file(file_path, show_errors=True, update_recent=True)

    def _save_to_file(self, file_path: str, *, show_errors: bool, update_recent: bool) -> bool:
        try:
            with open(file_path, 'w', encoding='utf-8') as file:
                file.write(self.editor.toPlainText())
            self.current_file = file_path
            self.editor.document().setModified(False)
            if update_recent:
                self._add_recent_file(file_path)
            return True
        except Exception as e:
            if show_errors:
                QMessageBox.critical(self, "Error", f"Could not save file: {str(e)}")
            return False

    def _autosave_tick(self) -> None:
        if not self.current_file:
            return
        if not self.editor.document().isModified():
            return

        self._save_to_file(self.current_file, show_errors=False, update_recent=False)

    def choose_preview_css(self) -> None:
        file_path, _ = QFileDialog.getOpenFileName(
            self, "Select Preview CSS", "", "CSS Files (*.css);;All Files (*)"
        )
        if not file_path:
            return
        self._custom_preview_css_path = file_path
        self._custom_preview_css_cache = ""
        self._custom_preview_css_cache_mtime = None
        self._settings.setValue("previewCssPath", self._custom_preview_css_path)
        self.update_preview()

    def clear_preview_css(self) -> None:
        self._custom_preview_css_path = ""
        self._custom_preview_css_cache = ""
        self._custom_preview_css_cache_mtime = None
        self._settings.setValue("previewCssPath", "")
        self.update_preview()

    def _get_custom_preview_css(self) -> str:
        path = (self._custom_preview_css_path or "").strip()
        if not path:
            return ""
        if not os.path.exists(path):

            self.clear_preview_css()
            return ""

        try:
            mtime = os.path.getmtime(path)
        except OSError:
            return ""

        if self._custom_preview_css_cache_mtime != mtime:
            try:
                with open(path, "r", encoding="utf-8") as f:
                    self._custom_preview_css_cache = f.read()
                self._custom_preview_css_cache_mtime = mtime
            except Exception:
                return ""

        return self._custom_preview_css_cache

    def _load_recent_files(self) -> list:
        value = self._settings.value("recentFiles", [])
        if value is None:
            return []
        if isinstance(value, str):
            return [value]
        if isinstance(value, (list, tuple)):
            return [str(v) for v in value if v]
        return []

    def _save_recent_files(self) -> None:
        self._settings.setValue("recentFiles", self._recent_files)

    def _add_recent_file(self, file_path: str) -> None:
        if not file_path:
            return
        file_path = os.path.abspath(file_path)
        self._recent_files = [p for p in self._recent_files if os.path.abspath(p) != file_path]
        self._recent_files.insert(0, file_path)
        self._recent_files = self._recent_files[:10]
        self._save_recent_files()
        self._rebuild_recent_files_menu()

    def _clear_recent_files(self) -> None:
        self._recent_files = []
        self._save_recent_files()
        self._rebuild_recent_files_menu()

    def _rebuild_recent_files_menu(self) -> None:
        if not hasattr(self, "recent_files_menu") or self.recent_files_menu is None:
            return

        self.recent_files_menu.clear()

        existing = [p for p in self._recent_files if p and os.path.exists(p)]
        self._recent_files = existing
        self._save_recent_files()

        if not self._recent_files:
            empty_action = self.recent_files_menu.addAction("(No recent files)")
            empty_action.setEnabled(False)
        else:
            for path in self._recent_files:
                action = self.recent_files_menu.addAction(path)
                action.triggered.connect(lambda checked=False, p=path: self.open_file_path(p))

        self.recent_files_menu.addSeparator()
        clear_action = self.recent_files_menu.addAction("Clear Recent Files")
        clear_action.setEnabled(bool(self._recent_files))
        clear_action.triggered.connect(self._clear_recent_files)

    def export_file(self, file_format):
        """Export the current markdown content to various formats."""

        file_filters = {
            'md': 'Markdown Files (*.md)',
            'txt': 'Text Files (*.txt)',
            'html': 'HTML Files (*.html)',
            'docx': 'Word Documents (*.docx)',
            'pdf': 'PDF Files (*.pdf)',
            'rtf': 'Rich Text Format (*.rtf)',
            'odt': 'OpenDocument Text (*.odt)'
        }

        if file_format not in file_filters:
            QMessageBox.warning(self, "Export Error", f"Unsupported export format: {file_format}")
            return


        if file_format == 'docx' and not DOCX_AVAILABLE:
            QMessageBox.warning(self, "Export Error", "python-docx library is not installed. Please install it with: pip install python-docx")
            return

        if file_format == 'pdf' and not (PDF_AVAILABLE or WEASYPRINT_AVAILABLE):
            QMessageBox.warning(self, "Export Error", "Neither reportlab nor weasyprint is installed. Please install one with: pip install reportlab or pip install weasyprint")
            return

        file_path, _ = QFileDialog.getSaveFileName(
            self, f"Export as {file_format.upper()}", "", file_filters[file_format]
        )

        if file_path:
            try:
                if file_format == 'md':
                    self.export_as_markdown(file_path)
                elif file_format == 'txt':
                    self.export_as_text(file_path)
                elif file_format == 'html':
                    self.export_as_html(file_path)
                elif file_format == 'docx':
                    self.export_as_docx(file_path)
                elif file_format == 'pdf':
                    self.export_as_pdf(file_path)
                elif file_format == 'rtf':
                    self.export_as_rtf(file_path)
                elif file_format == 'odt':
                    self.export_as_odt(file_path)

                QMessageBox.information(self, "Export Successful", f"File exported successfully to {file_path}")
            except Exception as e:
                QMessageBox.critical(self, "Export Error", f"Could not export file: {str(e)}")

    def export_as_markdown(self, file_path):

        with open(file_path, 'w', encoding='utf-8') as file:
            file.write(self.editor.toPlainText())

    def export_as_text(self, file_path):

        markdown_text = self.editor.toPlainText()

        import re
        text = markdown_text

        text = re.sub(r'^#{1,6}\s+', '', text, flags=re.MULTILINE)

        text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
        text = re.sub(r'\*(.*?)\*', r'\1', text)

        text = re.sub(r'```.*?\n(.*?)\n```', r'\1', text, flags=re.DOTALL)
        text = re.sub(r'`(.*?)`', r'\1', text)

        text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)

        text = re.sub(r'!\[([^\]]*)\]\([^\)]+\)', r'[\1]', text)

        with open(file_path, 'w', encoding='utf-8') as file:
            file.write(text)

    def export_as_html(self, file_path):

        markdown_text = self.editor.toPlainText()
        html = markdown.markdown(markdown_text, extensions=['codehilite', 'tables', 'toc'])

        pygments_css = ""
        if PYGMENTS_AVAILABLE:
            pygments_css = HtmlFormatter(style="default").get_style_defs('.codehilite')

        styled_html = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>Exported Markdown Document</title>
    <style>
        {pygments_css}
        body {{
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
            line-height: 1.6;
            color: #333;
            max-width: 800px;
            margin: 0 auto;
            padding: 20px;
            background-color: #fff;
        }}
        h1, h2, h3, h4, h5, h6 {{
            margin-top: 24px;
            margin-bottom: 16px;
            font-weight: 600;
            line-height: 1.25;
        }}
        h1 {{ font-size: 2em; border-bottom: 1px solid #eaecef; padding-bottom: 0.3em; }}
        h2 {{ font-size: 1.5em; border-bottom: 1px solid #eaecef; padding-bottom: 0.3em; }}
        h3 {{ font-size: 1.25em; }}
        h4 {{ font-size: 1em; }}
        h5 {{ font-size: 0.875em; }}
        h6 {{ font-size: 0.85em; color: #6a737d; }}
        p {{ margin-bottom: 16px; }}
        code {{
            background-color: #f6f8fa;
            border-radius: 3px;
            font-size: 85%;
            margin: 0;
            padding: 0.2em 0.4em;
        }}
        pre {{
            background-color: #f6f8fa;
            border-radius: 6px;
            padding: 16px;
            overflow: auto;
            font-size: 85%;
            line-height: 1.45;
        }}
        pre code {{
            background-color: transparent;
            border: 0;
            display: inline;
            line-height: inherit;
            margin: 0;
            max-width: auto;
            overflow: visible;
            padding: 0;
            word-wrap: normal;
        }}
        blockquote {{
            border-left: 0.25em solid #dfe2e5;
            color: #6a737d;
            padding: 0 1em;
            margin: 0 0 16px 0;
        }}
        table {{
            border-spacing: 0;
            border-collapse: collapse;
            margin-bottom: 16px;
        }}
        table th, table td {{
            border: 1px solid #dfe2e5;
            padding: 6px 13px;
        }}
        table th {{
            background-color: #f6f8fa;
            font-weight: 600;
        }}
        table tr:nth-child(2n) {{
            background-color: #f6f8fa;
        }}
        ul, ol {{
            padding-left: 2em;
            margin-bottom: 16px;
        }}
        li {{
            margin-bottom: 0.25em;
        }}
        a {{
            color: #0366d6;
            text-decoration: none;
        }}
        a:hover {{
            text-decoration: underline;
        }}
        img {{
            max-width: 100%;
            height: auto;
        }}
        hr {{
            border: none;
            border-top: 1px solid #eaecef;
            height: 1px;
            margin: 24px 0;
        }}
    </style>
</head>
<body>
    {html}
</body>
</html>"""

        with open(file_path, 'w', encoding='utf-8') as file:
            file.write(styled_html)

    def export_as_docx(self, file_path):

        if not DOCX_AVAILABLE:
            raise ImportError("python-docx library is not available")

        markdown_text = self.editor.toPlainText()
        html = markdown.markdown(markdown_text, extensions=['codehilite', 'tables', 'toc'])


        doc = Document()


        doc.add_heading('Exported Markdown Document', 0)


        lines = markdown_text.split('\n')
        in_code_block = False
        code_content = []

        for line in lines:
            if line.strip().startswith('```'):
                if in_code_block:

                    if code_content:
                        paragraph = doc.add_paragraph()
                        run = paragraph.add_run('\n'.join(code_content))
                        run.font.name = 'Courier New'
                        code_content = []
                    in_code_block = False
                else:

                    in_code_block = True
                continue

            if in_code_block:
                code_content.append(line)
                continue


            if line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.startswith('#### '):
                doc.add_heading(line[5:], level=4)
            elif line.startswith('##### '):
                doc.add_heading(line[6:], level=5)
            elif line.startswith('###### '):
                doc.add_heading(line[7:], level=6)

            elif line.strip() == '---' or line.strip() == '***':
                doc.add_paragraph('_' * 50)

            elif line.strip().startswith(('- ', '* ', '+ ')):
                p = doc.add_paragraph(line.strip()[2:], style='List Bullet')
            elif line.strip() and line.strip()[0].isdigit() and line.strip().find('.') == 1:
                p = doc.add_paragraph(line.strip()[3:], style='List Number')

            elif not line.strip():
                doc.add_paragraph()

            elif line.strip():

                processed_line = line

                processed_line = re.sub(r'\*\*(.*?)\*\*', r'\1', processed_line)

                processed_line = re.sub(r'\*(.*?)\*', r'\1', processed_line)

                processed_line = re.sub(r'`(.*?)`', r'\1', processed_line)

                processed_line = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', processed_line)

                doc.add_paragraph(processed_line)

        doc.save(file_path)

    def export_as_pdf(self, file_path):

        markdown_text = self.editor.toPlainText()


        if WEASYPRINT_AVAILABLE:
            html = markdown.markdown(markdown_text, extensions=['codehilite', 'tables', 'toc'])

            styled_html = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>Exported Markdown Document</title>
    <style>
        @page {{
            size: letter;
            margin: 1in;
        }}
        body {{
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
            line-height: 1.6;
            color: #333;
            font-size: 12pt;
        }}
        h1, h2, h3, h4, h5, h6 {{
            margin-top: 24px;
            margin-bottom: 16px;
            font-weight: 600;
            line-height: 1.25;
        }}
        h1 {{ font-size: 2em; border-bottom: 1px solid #eaecef; padding-bottom: 0.3em; }}
        h2 {{ font-size: 1.5em; border-bottom: 1px solid #eaecef; padding-bottom: 0.3em; }}
        h3 {{ font-size: 1.25em; }}
        h4 {{ font-size: 1em; }}
        h5 {{ font-size: 0.875em; }}
        h6 {{ font-size: 0.85em; color: #6a737d; }}
        p {{ margin-bottom: 16px; }}
        code {{
            background-color: #f6f8fa;
            border-radius: 3px;
            font-size: 85%;
            margin: 0;
            padding: 0.2em 0.4em;
        }}
        pre {{
            background-color: #f6f8fa;
            border-radius: 6px;
            padding: 16px;
            overflow: auto;
            font-size: 85%;
            line-height: 1.45;
        }}
        pre code {{
            background-color: transparent;
            border: 0;
            display: inline;
            line-height: inherit;
            margin: 0;
            max-width: auto;
            overflow: visible;
            padding: 0;
            word-wrap: normal;
        }}
        blockquote {{
            border-left: 0.25em solid #dfe2e5;
            color: #6a737d;
            padding: 0 1em;
            margin: 0 0 16px 0;
        }}
        table {{
            border-spacing: 0;
            border-collapse: collapse;
            margin-bottom: 16px;
        }}
        table th, table td {{
            border: 1px solid #dfe2e5;
            padding: 6px 13px;
        }}
        table th {{
            background-color: #f6f8fa;
            font-weight: 600;
        }}
        table tr:nth-child(2n) {{
            background-color: #f6f8fa;
        }}
        ul, ol {{
            padding-left: 2em;
            margin-bottom: 16px;
        }}
        li {{
            margin-bottom: 0.25em;
        }}
        a {{
            color: #0366d6;
            text-decoration: none;
        }}
        img {{
            max-width: 100%;
            height: auto;
        }}
        hr {{
            border: none;
            border-top: 1px solid #eaecef;
            height: 1px;
            margin: 24px 0;
        }}
    </style>
</head>
<body>
    {html}
</body>
</html>"""


            weasyprint.HTML(string=styled_html).write_pdf(file_path)


        elif PDF_AVAILABLE:
            doc = SimpleDocTemplate(file_path, pagesize=letter)
            styles = getSampleStyleSheet()
            story = []


            title_style = styles['Title']
            title = Paragraph("Exported Markdown Document", title_style)
            story.append(title)
            story.append(Spacer(1, 12))


            lines = markdown_text.split('\n')
            in_code_block = False
            code_content = []

            for line in lines:
                if line.strip().startswith('```'):
                    if in_code_block:

                        if code_content:
                            code_text = '\n'.join(code_content)
                            code_style = styles['Code']
                            code_para = Paragraph(code_text, code_style)
                            story.append(code_para)
                            story.append(Spacer(1, 6))
                            code_content = []
                        in_code_block = False
                    else:

                        in_code_block = True
                    continue

                if in_code_block:
                    code_content.append(line)
                    continue


                if line.startswith('# '):
                    story.append(Paragraph(line[2:], styles['Heading1']))
                elif line.startswith('## '):
                    story.append(Paragraph(line[3:], styles['Heading2']))
                elif line.startswith('### '):
                    story.append(Paragraph(line[4:], styles['Heading3']))
                elif line.startswith('#### '):
                    story.append(Paragraph(line[5:], styles['Heading4']))
                elif line.startswith('##### '):
                    story.append(Paragraph(line[6:], styles['Heading5']))
                elif line.startswith('###### '):
                    story.append(Paragraph(line[7:], styles['Heading6']))

                elif line.strip() == '---' or line.strip() == '***':
                    story.append(Spacer(1, 12))

                elif line.strip().startswith(('- ', '* ', '+ ')):
                    story.append(Paragraph(f"• {line.strip()[2:]}", styles['Normal']))
                elif line.strip() and line.strip()[0].isdigit() and line.strip().find('.') == 1:
                    story.append(Paragraph(f"{line.strip()}", styles['Normal']))

                elif not line.strip():
                    story.append(Spacer(1, 6))

                elif line.strip():

                    processed_line = line

                    processed_line = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', processed_line)

                    processed_line = re.sub(r'\*(.*?)\*', r'<i>\1</i>', processed_line)

                    processed_line = re.sub(r'`(.*?)`', r'<font name="Courier">\1</font>', processed_line)

                    processed_line = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', processed_line)

                    story.append(Paragraph(processed_line, styles['Normal']))

            doc.build(story)

        else:
            raise ImportError("Neither weasyprint nor reportlab is available")

    def export_as_rtf(self, file_path):

        markdown_text = self.editor.toPlainText()
        html = markdown.markdown(markdown_text, extensions=['codehilite', 'tables', 'toc'])


        rtf_content = r"{\rtf1\ansi\deff0"
        rtf_content += r"{\fonttbl{\f0 Times New Roman;}}"
        rtf_content += r"{\colortbl;\red0\green0\blue0;}"
        rtf_content += r"\fs24"


        lines = markdown_text.split('\n')
        in_code_block = False
        code_content = []

        for line in lines:
            if line.strip().startswith('```'):
                if in_code_block:

                    if code_content:
                        rtf_content += r"{\pard\plain\f0\fs20 "
                        for code_line in code_content:
                            rtf_content += code_line.replace('\\', '\\\\').replace('{', '\\{').replace('}', '\\}') + r"\line "
                        rtf_content += r"}\par"
                        code_content = []
                    in_code_block = False
                else:

                    in_code_block = True
                continue

            if in_code_block:
                code_content.append(line)
                continue


            if line.startswith('# '):
                rtf_content += r"{\pard\plain\f0\fs36\b " + line[2:].replace('\\', '\\\\').replace('{', '\\{').replace('}', '\\}') + r"}\par"
            elif line.startswith('## '):
                rtf_content += r"{\pard\plain\f0\fs32\b " + line[3:].replace('\\', '\\\\').replace('{', '\\{').replace('}', '\\}') + r"}\par"
            elif line.startswith('### '):
                rtf_content += r"{\pard\plain\f0\fs28\b " + line[4:].replace('\\', '\\\\').replace('{', '\\{').replace('}', '\\}') + r"}\par"
            elif line.startswith('#### '):
                rtf_content += r"{\pard\plain\f0\fs26\b " + line[5:].replace('\\', '\\\\').replace('{', '\\{').replace('}', '\\}') + r"}\par"
            elif line.startswith('##### '):
                rtf_content += r"{\pard\plain\f0\fs24\b " + line[6:].replace('\\', '\\\\').replace('{', '\\{').replace('}', '\\}') + r"}\par"
            elif line.startswith('###### '):
                rtf_content += r"{\pard\plain\f0\fs22\b " + line[7:].replace('\\', '\\\\').replace('{', '\\{').replace('}', '\\}') + r"}\par"

            elif line.strip() == '---' or line.strip() == '***':
                rtf_content += r"{\pard\plain\f0\fs24 " + "_" * 50 + r"}\par"

            elif line.strip().startswith(('- ', '* ', '+ ')):
                rtf_content += r"{\pard\plain\f0\fs24 \bullet " + line.strip()[2:].replace('\\', '\\\\').replace('{', '\\{').replace('}', '\\}') + r"}\par"
            elif line.strip() and line.strip()[0].isdigit() and line.strip().find('.') == 1:
                rtf_content += r"{\pard\plain\f0\fs24 " + line.strip().replace('\\', '\\\\').replace('{', '\\{').replace('}', '\\}') + r"}\par"

            elif not line.strip():
                rtf_content += r"\par"

            elif line.strip():

                processed_line = line

                processed_line = re.sub(r'\*\*(.*?)\*\*', r'\b \1\b0', processed_line)

                processed_line = re.sub(r'\*(.*?)\*', r'\i \1\i0', processed_line)

                processed_line = re.sub(r'`(.*?)`', r'\f1\fs18 \1\f0\fs24', processed_line)

                processed_line = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', processed_line)

                rtf_content += r"{\pard\plain\f0\fs24 " + processed_line.replace('\\', '\\\\').replace('{', '\\{').replace('}', '\\}') + r"}\par"

        rtf_content += "}"

        with open(file_path, 'w', encoding='utf-8') as file:
            file.write(rtf_content)

    def export_as_odt(self, file_path):


        import zipfile
        import xml.etree.ElementTree as ET

        markdown_text = self.editor.toPlainText()
        html = markdown.markdown(markdown_text, extensions=['codehilite', 'tables', 'toc'])


        content = ET.Element("office:document-content")
        content.set("xmlns:office", "urn:oasis:names:tc:opendocument:xmlns:office:1.0")
        content.set("xmlns:text", "urn:oasis:names:tc:opendocument:xmlns:text:1.0")
        content.set("xmlns:style", "urn:oasis:names:tc:opendocument:xmlns:style:1.0")
        content.set("office:version", "1.0")

        body = ET.SubElement(content, "office:body")
        text = ET.SubElement(body, "office:text")


        lines = markdown_text.split('\n')
        in_code_block = False
        code_content = []

        for line in lines:
            if line.strip().startswith('```'):
                if in_code_block:

                    if code_content:
                        for code_line in code_content:
                            p = ET.SubElement(text, "text:p")
                            p.text = code_line
                        code_content = []
                    in_code_block = False
                else:

                    in_code_block = True
                continue

            if in_code_block:
                code_content.append(line)
                continue


            if line.startswith('# '):
                h = ET.SubElement(text, "text:h", attrib={"text:outline-level": "1"})
                h.text = line[2:]
            elif line.startswith('## '):
                h = ET.SubElement(text, "text:h", attrib={"text:outline-level": "2"})
                h.text = line[3:]
            elif line.startswith('### '):
                h = ET.SubElement(text, "text:h", attrib={"text:outline-level": "3"})
                h.text = line[4:]
            elif line.startswith('#### '):
                h = ET.SubElement(text, "text:h", attrib={"text:outline-level": "4"})
                h.text = line[5:]
            elif line.startswith('##### '):
                h = ET.SubElement(text, "text:h", attrib={"text:outline-level": "5"})
                h.text = line[6:]
            elif line.startswith('###### '):
                h = ET.SubElement(text, "text:h", attrib={"text:outline-level": "6"})
                h.text = line[7:]

            elif line.strip() == '---' or line.strip() == '***':
                p = ET.SubElement(text, "text:p")
                p.text = "_" * 50

            elif line.strip().startswith(('- ', '* ', '+ ')):
                p = ET.SubElement(text, "text:p")
                p.text = "• " + line.strip()[2:]
            elif line.strip() and line.strip()[0].isdigit() and line.strip().find('.') == 1:
                p = ET.SubElement(text, "text:p")
                p.text = line.strip()

            elif not line.strip():
                p = ET.SubElement(text, "text:p")
                p.text = ""

            elif line.strip():

                processed_line = line

                processed_line = re.sub(r'\*\*(.*?)\*\*', r'\1', processed_line)

                processed_line = re.sub(r'\*(.*?)\*', r'\1', processed_line)

                processed_line = re.sub(r'`(.*?)`', r'\1', processed_line)

                processed_line = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', processed_line)

                p = ET.SubElement(text, "text:p")
                p.text = processed_line


        with zipfile.ZipFile(file_path, 'w', zipfile.ZIP_DEFLATED) as odt:

            odt.writestr('mimetype', 'application/vnd.oasis.opendocument.text')


            content_str = ET.tostring(content, encoding='unicode', xml_declaration=True)
            odt.writestr('content.xml', content_str)


            manifest = ET.Element("manifest:manifest")
            manifest.set("xmlns:manifest", "urn:oasis:names:tc:opendocument:xmlns:manifest:1.0")

            file_entry1 = ET.SubElement(manifest, "manifest:file-entry")
            file_entry1.set("manifest:full-path", "/")
            file_entry1.set("manifest:media-type", "application/vnd.oasis.opendocument.text")

            file_entry2 = ET.SubElement(manifest, "manifest:file-entry")
            file_entry2.set("manifest:full-path", "content.xml")
            file_entry2.set("manifest:media-type", "text/xml")

            manifest_str = ET.tostring(manifest, encoding='unicode', xml_declaration=True)
            odt.writestr('META-INF/manifest.xml', manifest_str)


def main():
    app = QApplication(sys.argv)
    editor = MarkdownEditor()
    editor.show()
    sys.exit(app.exec())


if __name__ == "__main__":
    main()
