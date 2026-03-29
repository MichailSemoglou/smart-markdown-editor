"""Built-in concrete exporters for Smart Markdown Editor.

Each class in this module extends :class:`src.exporters.BaseExporter` and
implements one output format.  Call :func:`register_all` once at startup to
make all exporters available through the registry in :mod:`src.exporters`.

Shared inline helpers (:func:`_parse_markdown_lines`, :func:`_strip_inline_md`)
avoid repeating the same parsing logic across every exporter.
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import Generator

from src.exporters import BaseExporter, register_exporter


# =============================================================================
# Shared parsing utilities
# =============================================================================

def _strip_inline_md(text: str) -> str:
    """Remove inline markdown markers from *text* and return plain content."""
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"__(.*?)__", r"\1", text)
    text = re.sub(r"\*(.*?)\*", r"\1", text)
    text = re.sub(r"_(.*?)_", r"\1", text)
    text = re.sub(r"`(.*?)`", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\([^\)]+\)", r"\1", text)
    text = re.sub(r"!\[([^\]]*)\]\([^\)]+\)", r"[\1]", text)
    return text


def _parse_markdown_lines(
    text: str,
) -> Generator[tuple[str, object], None, None]:
    """Yield ``(block_type, content)`` tuples for each logical block.

    block_type values
    -----------------
    ``'h1'`` .. ``'h6'``  — heading at that level; *content* is the heading text
    ``'code_block'``       — fenced code block; *content* is a ``list[str]``
    ``'hr'``               — horizontal rule; *content* is ``''``
    ``'list_bullet'``      — unordered list item; *content* is the item text
    ``'list_ordered'``     — ordered list item; *content* is the item text
    ``'empty'``            — blank line; *content* is ``''``
    ``'paragraph'``        — body text; *content* is the raw line
    """
    lines = text.split("\n")
    in_code_block = False
    code_lines: list[str] = []
    for line in lines:
        if line.strip().startswith("```"):
            if in_code_block:
                yield ("code_block", code_lines)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            continue
        if in_code_block:
            code_lines.append(line)
            continue
        heading_match = re.match(r"^(#{1,6})\s+(.*)", line)
        if heading_match:
            level = len(heading_match.group(1))
            yield (f"h{level}", heading_match.group(2))
            continue
        if line.strip() in ("---", "***", "___"):
            yield ("hr", "")
            continue
        if re.match(r"^\s*[-*+]\s+", line):
            yield ("list_bullet", line.strip()[2:].strip())
            continue
        if re.match(r"^\s*\d+\.\s+", line):
            m = re.match(r"^\s*\d+\.\s+(.*)", line)
            yield ("list_ordered", m.group(1) if m else line)
            continue
        if not line.strip():
            yield ("empty", "")
            continue
        yield ("paragraph", line)
    # flush an unclosed fence
    if in_code_block and code_lines:
        yield ("code_block", code_lines)


# =============================================================================
# Concrete exporters
# =============================================================================


class MarkdownExporter(BaseExporter):
    """Pass-through exporter: saves markdown as-is."""

    name = "Markdown"
    extension = "md"
    file_filter = "Markdown Files (*.md)"

    def export(self, content: str, output_path: Path) -> None:
        output_path.write_text(content, encoding="utf-8")


class TextExporter(BaseExporter):
    """Strip markdown syntax and export as plain text."""

    name = "Plain Text"
    extension = "txt"
    file_filter = "Text Files (*.txt)"

    def export(self, content: str, output_path: Path) -> None:
        text = content
        text = re.sub(r"^#{1,6}\s+", "", text, flags=re.MULTILINE)
        text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
        text = re.sub(r"\*(.*?)\*", r"\1", text)
        text = re.sub(r"```.*?\n(.*?)\n```", r"\1", text, flags=re.DOTALL)
        text = re.sub(r"`(.*?)`", r"\1", text)
        text = re.sub(r"\[([^\]]+)\]\([^\)]+\)", r"\1", text)
        text = re.sub(r"!\[([^\]]*)\]\([^\)]+\)", r"[\1]", text)
        output_path.write_text(text, encoding="utf-8")


class HTMLExporter(BaseExporter):
    """Export as a standalone styled HTML document."""

    name = "HTML"
    extension = "html"
    file_filter = "HTML Files (*.html)"

    def export(self, content: str, output_path: Path) -> None:
        import markdown as md_lib

        try:
            from pygments.formatters import HtmlFormatter
            pygments_css = HtmlFormatter(style="default").get_style_defs(".codehilite")
        except ImportError:
            pygments_css = ""

        html_body = md_lib.markdown(
            content, extensions=["codehilite", "tables", "toc"]
        )

        styled = f"""<!DOCTYPE html>
<html>
<head>
  <meta charset="UTF-8">
  <title>Exported Markdown</title>
  <style>
    {pygments_css}
    body {{
      font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
      line-height: 1.6; color: #333; max-width: 800px;
      margin: 0 auto; padding: 20px; background-color: #fff;
    }}
    h1, h2, h3, h4, h5, h6 {{
      margin-top: 24px; margin-bottom: 16px; font-weight: 600;
    }}
    h1 {{ font-size: 2em; border-bottom: 1px solid #eaecef; padding-bottom: 0.3em; }}
    h2 {{ font-size: 1.5em; border-bottom: 1px solid #eaecef; padding-bottom: 0.3em; }}
    code {{ background-color: #f6f8fa; border-radius: 3px; padding: 0.2em 0.4em; }}
    pre {{ background-color: #f6f8fa; border-radius: 6px; padding: 16px; overflow: auto; }}
    .codehilite {{ background-color: #f6f8fa; border-radius: 6px; padding: 16px; }}
    .codehilite pre {{ margin: 0; padding: 0; background: transparent; }}
    blockquote {{ border-left: 0.25em solid #eaecef; color: #6a737d; padding: 0 1em; }}
    table {{ border-collapse: collapse; margin-bottom: 16px; }}
    table th, table td {{ border: 1px solid #eaecef; padding: 6px 13px; }}
    table th {{ background-color: #f6f8fa; font-weight: 600; }}
    a {{ color: #0366d6; text-decoration: none; }}
    img {{ max-width: 100%; height: auto; }}
  </style>
</head>
<body>
{html_body}
</body>
</html>"""
        output_path.write_text(styled, encoding="utf-8")


class DocxExporter(BaseExporter):
    """Export as a Microsoft Word (.docx) document."""

    name = "Word Document"
    extension = "docx"
    file_filter = "Word Documents (*.docx)"
    requires_library = "docx"

    def export(self, content: str, output_path: Path) -> None:
        from docx import Document
        from docx.shared import Pt

        doc = Document()
        for btype, bdata in _parse_markdown_lines(content):
            if btype.startswith("h") and btype != "hr":
                level = int(btype[1])
                doc.add_heading(_strip_inline_md(str(bdata)), level=level)
            elif btype == "code_block":
                for code_line in bdata:  # type: ignore[union-attr, attr-defined]
                    para = doc.add_paragraph(code_line)
                    para.style = "No Spacing"
                    run = para.runs[0] if para.runs else para.add_run()
                    run.font.name = "Courier New"
                    run.font.size = Pt(10)
            elif btype == "hr":
                doc.add_paragraph("─" * 40)
            elif btype in ("list_bullet", "list_ordered"):
                style = (
                    "List Bullet" if btype == "list_bullet" else "List Number"
                )
                try:
                    doc.add_paragraph(_strip_inline_md(str(bdata)), style=style)
                except KeyError:
                    doc.add_paragraph(_strip_inline_md(str(bdata)))
            elif btype == "empty":
                doc.add_paragraph("")
            elif btype == "paragraph":
                doc.add_paragraph(_strip_inline_md(str(bdata)))
        doc.save(str(output_path))


class PDFExporter(BaseExporter):
    """Export as PDF using weasyprint (preferred) or reportlab."""

    name = "PDF Document"
    extension = "pdf"
    file_filter = "PDF Files (*.pdf)"

    @property
    def is_available(self) -> bool:
        try:
            import weasyprint  # noqa: F401
            return True
        except ImportError:
            pass
        try:
            import reportlab  # noqa: F401
            return True
        except ImportError:
            return False

    def export(self, content: str, output_path: Path) -> None:
        try:
            self._export_weasyprint(content, output_path)
        except ImportError:
            self._export_reportlab(content, output_path)

    def _export_weasyprint(self, content: str, output_path: Path) -> None:
        import markdown as md_lib
        import weasyprint

        html_body = md_lib.markdown(content, extensions=["tables", "toc"])
        html = f"""<!DOCTYPE html><html><head><meta charset="UTF-8">
<style>
  body {{ font-family: sans-serif; line-height: 1.6; max-width: 800px;
         margin: 0 auto; padding: 20px; }}
  h1, h2, h3, h4, h5, h6 {{ margin-top: 1em; }}
  code, pre {{ background: #f6f8fa; border-radius: 3px; padding: 0.2em 0.4em; }}
  pre {{ padding: 1em; }}
  table {{ border-collapse: collapse; width: 100%; }}
  table th, table td {{ border: 1px solid #ccc; padding: 6px 10px; }}
</style>
</head><body>{html_body}</body></html>"""
        weasyprint.HTML(string=html).write_pdf(str(output_path))

    def _export_reportlab(self, content: str, output_path: Path) -> None:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.lib.units import inch
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

        doc = SimpleDocTemplate(str(output_path), pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        for btype, bdata in _parse_markdown_lines(content):
            if btype.startswith("h") and btype != "hr":
                level = int(btype[1])
                style_name = f"Heading{min(level, 6)}"
                story.append(
                    Paragraph(_strip_inline_md(str(bdata)), styles[style_name])
                )
                story.append(Spacer(1, 0.1 * inch))
            elif btype == "code_block":
                for code_line in bdata:  # type: ignore[union-attr, attr-defined]
                    story.append(Paragraph(code_line, styles["Code"]))
            elif btype == "hr":
                story.append(Spacer(1, 0.2 * inch))
            elif btype in ("list_bullet", "list_ordered"):
                bullet = "• " if btype == "list_bullet" else ""
                story.append(
                    Paragraph(
                        bullet + _strip_inline_md(str(bdata)), styles["Normal"]
                    )
                )
            elif btype == "paragraph":
                story.append(
                    Paragraph(_strip_inline_md(str(bdata)), styles["Normal"])
                )
                story.append(Spacer(1, 0.1 * inch))
            elif btype == "empty":
                story.append(Spacer(1, 0.1 * inch))
        doc.build(story)


class RTFExporter(BaseExporter):
    """Export as Rich Text Format (.rtf) — pure-Python, no extra library."""

    name = "Rich Text Format"
    extension = "rtf"
    file_filter = "Rich Text Format (*.rtf)"

    def export(self, content: str, output_path: Path) -> None:
        lines = [
            r"{\rtf1\ansi\deff0",
            r"{\fonttbl{\f0 Times New Roman;}{\f1 Courier New;}}",
            r"{\colortbl;\red0\green0\blue0;\red68\green68\blue68;}",
            r"\f0\fs24",
        ]
        for btype, bdata in _parse_markdown_lines(content):
            if btype.startswith("h") and btype != "hr":
                level = int(btype[1])
                size = max(28 - (level - 1) * 4, 20)
                lines.append(
                    rf"\pard\sb200\sa100\b\fs{size} "
                    + self._rtf_escape(_strip_inline_md(str(bdata)))
                    + r"\b0\par"
                )
            elif btype == "code_block":
                lines.append(r"\pard\f1\fs18")
                for code_line in bdata:  # type: ignore[union-attr, attr-defined]
                    lines.append(self._rtf_escape(code_line) + r"\line")
                lines.append(r"\f0\fs24\par")
            elif btype == "hr":
                lines.append(r"\pard\brdrb\brdrs\brdrw10 \par")
            elif btype in ("list_bullet", "list_ordered"):
                bullet = "\u2022  " if btype == "list_bullet" else ""
                lines.append(
                    r"\pard\li360 "
                    + self._rtf_escape(bullet + _strip_inline_md(str(bdata)))
                    + r"\par"
                )
            elif btype == "empty":
                lines.append(r"\par")
            elif btype == "paragraph":
                lines.append(
                    r"\pard\sb0\sa100 "
                    + self._rtf_escape(_strip_inline_md(str(bdata)))
                    + r"\par"
                )
        lines.append("}")
        output_path.write_text("\n".join(lines), encoding="ascii", errors="replace")

    @staticmethod
    def _rtf_escape(text: str) -> str:
        """Escape special RTF characters and encode non-ASCII as \\uN?."""
        text = text.replace("\\", "\\\\").replace("{", "\\{").replace("}", "\\}")
        result = []
        for ch in text:
            code = ord(ch)
            if code <= 127:
                result.append(ch)
            else:
                result.append(f"\\u{code}?")
        return "".join(result)


class ODTExporter(BaseExporter):
    """Export as OpenDocument Text (.odt) — pure-Python, no extra library."""

    name = "OpenDocument Text"
    extension = "odt"
    file_filter = "OpenDocument Text (*.odt)"

    def export(self, content: str, output_path: Path) -> None:
        import xml.etree.ElementTree as ET
        import zipfile

        root = ET.Element("office:document-content")
        root.set("xmlns:office", "urn:oasis:names:tc:opendocument:xmlns:office:1.0")
        root.set("xmlns:text", "urn:oasis:names:tc:opendocument:xmlns:text:1.0")
        root.set("xmlns:style", "urn:oasis:names:tc:opendocument:xmlns:style:1.0")
        root.set("office:version", "1.0")

        body = ET.SubElement(root, "office:body")
        text_el = ET.SubElement(body, "office:text")

        for btype, bdata in _parse_markdown_lines(content):
            if btype.startswith("h") and btype != "hr":
                level = btype[1:]
                h = ET.SubElement(
                    text_el, "text:h", attrib={"text:outline-level": level}
                )
                h.text = _strip_inline_md(str(bdata))
            elif btype == "code_block":
                for code_line in bdata:  # type: ignore[union-attr, attr-defined]
                    p = ET.SubElement(text_el, "text:p")
                    p.text = code_line
            elif btype == "hr":
                p = ET.SubElement(text_el, "text:p")
                p.text = "_" * 50
            elif btype in ("list_bullet", "list_ordered"):
                p = ET.SubElement(text_el, "text:p")
                prefix = "\u2022 " if btype == "list_bullet" else ""
                p.text = prefix + _strip_inline_md(str(bdata))
            elif btype == "empty":
                ET.SubElement(text_el, "text:p")
            elif btype == "paragraph":
                p = ET.SubElement(text_el, "text:p")
                p.text = _strip_inline_md(str(bdata))

        content_xml = ET.tostring(root, encoding="unicode", xml_declaration=True)

        manifest_root = ET.Element("manifest:manifest")
        manifest_root.set(
            "xmlns:manifest",
            "urn:oasis:names:tc:opendocument:xmlns:manifest:1.0",
        )
        e1 = ET.SubElement(manifest_root, "manifest:file-entry")
        e1.set("manifest:full-path", "/")
        e1.set("manifest:media-type", "application/vnd.oasis.opendocument.text")
        e2 = ET.SubElement(manifest_root, "manifest:file-entry")
        e2.set("manifest:full-path", "content.xml")
        e2.set("manifest:media-type", "text/xml")
        manifest_xml = ET.tostring(
            manifest_root, encoding="unicode", xml_declaration=True
        )

        with zipfile.ZipFile(str(output_path), "w", zipfile.ZIP_DEFLATED) as odt:
            odt.writestr("mimetype", "application/vnd.oasis.opendocument.text")
            odt.writestr("content.xml", content_xml)
            odt.writestr("META-INF/manifest.xml", manifest_xml)


# =============================================================================
# Registration
# =============================================================================


def register_all() -> None:
    """Register every built-in exporter in the global registry."""
    for cls in (
        MarkdownExporter,
        TextExporter,
        HTMLExporter,
        DocxExporter,
        PDFExporter,
        RTFExporter,
        ODTExporter,
    ):
        register_exporter(cls)
